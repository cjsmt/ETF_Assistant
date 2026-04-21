"""
从 trace + Agent 回复生成带图、表、分析、信息架构的 HTML 周报。
支持导出为 HTML、Word(.docx)、PDF。
用法: python scripts/generate_report_html.py [trace路径] [--format html|docx|pdf|all]
"""
import argparse
import base64
import io
import os
import sys
import tempfile
from html import escape

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

if __package__ in {None, ""}:
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.report_data import build_report_data, load_agent_response, load_trace

plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False


def _fig_to_base64(fig) -> str:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=100, bbox_inches="tight")
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


def _to_float_weight(value) -> float:
    try:
        return float(str(value).replace("%", "").strip())
    except (TypeError, ValueError):
        return 0.0


def _make_quadrant_chart(data: dict) -> str:
    """四象限行业数量柱状图"""
    quadrant = data.get("quadrant", {})
    labels = ["黄金配置区", "左侧观察区", "高危警示区", "垃圾规避区"]
    counts = [
        len(quadrant.get("golden", [])),
        len(quadrant.get("left", [])),
        len(quadrant.get("danger", [])),
        len(quadrant.get("garbage", [])),
    ]
    colors = ["#2ecc71", "#3498db", "#f39c12", "#e74c3c"]

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(labels, counts, color=colors, edgecolor="white", linewidth=1.2)
    ax.set_ylabel("行业数量")
    ax.set_title("四象限行业分布")
    for bar in bars:
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.2,
            str(int(bar.get_height())),
            ha="center",
            fontsize=11,
        )
    plt.tight_layout()
    encoded = _fig_to_base64(fig)
    plt.close()
    return encoded


def _make_weights_pie(data: dict) -> str:
    """组合权重饼图"""
    items = []
    for row in data.get("etf_rows", []):
        weight = _to_float_weight(row.get("weight"))
        if weight > 0:
            items.append((row.get("sector", ""), weight))

    cash_weight = _to_float_weight(data.get("cash_reserve", "0"))
    if cash_weight > 0:
        items.append(("现金", cash_weight))

    if not items:
        fig, ax = plt.subplots(figsize=(5, 5))
        ax.text(0.5, 0.5, "无数据", ha="center", va="center", fontsize=14)
        encoded = _fig_to_base64(fig)
        plt.close()
        return encoded

    labels = [item[0] for item in items]
    sizes = [item[1] for item in items]
    colors = plt.cm.Set3([i / max(len(sizes), 1) for i in range(len(sizes))])

    fig, ax = plt.subplots(figsize=(7, 7))
    ax.pie(sizes, labels=labels, autopct="%1.0f%%", colors=colors, startangle=90)
    ax.set_title("ETF 组合权重")
    plt.tight_layout()
    encoded = _fig_to_base64(fig)
    plt.close()
    return encoded


def _build_html(data: dict, agent_response: str, chart_quadrant: str, chart_pie: str) -> str:
    risk = data.get("risk_checks", {})
    obs = data.get("observation_pool", {})
    news = data.get("news_validation", {})
    veto = data.get("veto_list", [])
    etf_rows = data.get("etf_rows", [])

    news_html = "".join(f"<li>{escape(k)}：{escape(str(v))}</li>" for k, v in (news or {}).items()) or "<li>无</li>"
    risk_macro = "".join(f"<li>{escape(r)}</li>" for r in risk.get("macro_risks", []))
    risk_sector = "".join(f"<li>{escape(r)}</li>" for r in risk.get("sector_risks", []))
    etf_rows_html = "".join(
        (
            f"<tr><td>{escape(row.get('sector', ''))}</td>"
            f"<td>{escape(str(row.get('code', '-')))}</td>"
            f"<td>{escape(str(row.get('weight', '-')))}</td>"
            f"<td>{escape(str(row.get('rationale', '-')))}</td></tr>"
        )
        for row in etf_rows
    )

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>ETF 行业轮动周报 · {escape(data.get('report_week', 'N/A'))}</title>
<style>
*{{box-sizing:border-box}}
body{{font-family:'Microsoft YaHei',sans-serif;margin:0;padding:24px;background:#f5f6fa;color:#2c3e50;line-height:1.6}}
.container{{max-width:900px;margin:0 auto}}
h1{{color:#1a252f;border-bottom:2px solid #3498db;padding-bottom:8px;font-size:1.8em}}
h2{{color:#2980b9;margin-top:32px;font-size:1.3em}}
.meta{{background:#fff;padding:16px;border-radius:8px;margin-bottom:24px;border-left:4px solid #3498db}}
.section{{background:#fff;padding:20px;margin-bottom:20px;border-radius:8px;box-shadow:0 1px 3px rgba(0,0,0,.08)}}
table{{width:100%;border-collapse:collapse;margin:12px 0}}
th,td{{border:1px solid #ddd;padding:10px;text-align:left}}
th{{background:#3498db;color:#fff}}
tr:nth-child(even){{background:#f8f9fa}}
.chart{{margin:20px 0;text-align:center}}
.chart img{{max-width:100%;height:auto;border-radius:4px}}
.agent-response{{background:#f8f9fa;padding:16px;border-radius:8px;white-space:pre-wrap;font-size:0.95em;border-left:4px solid #2ecc71}}
.reasoning{{background:#ecf0f1;padding:12px;border-radius:6px;font-style:italic;color:#34495e}}
.toc{{background:#fff;padding:16px;border-radius:8px;margin-bottom:24px}}
.toc ul{{list-style:none;padding:0}}
.toc li{{padding:6px 0;border-bottom:1px solid #eee}}
.toc a{{color:#3498db;text-decoration:none}}
.toc a:hover{{text-decoration:underline}}
</style>
</head>
<body>
<div class="container">
<h1>📊 ETF 行业轮动周报</h1>

<div class="meta">
<strong>报告周</strong>：{escape(data.get('report_week', 'N/A'))} &nbsp;|&nbsp;
<strong>报告日期</strong>：{escape(data.get('decision_date', 'N/A'))} &nbsp;|&nbsp;
<strong>信号截止</strong>：{escape(data.get('decision_date', 'N/A'))} &nbsp;|&nbsp;
<strong>因子计算区间</strong>：{escape(str(data.get('data_timestamp', '-')))} &nbsp;|&nbsp;
<strong>配置版本</strong>：{escape(str(data.get('config_version', 'N/A')))}
</div>

<nav class="toc">
<strong>目录</strong>
<ul>
<li><a href="#quadrant">一、四象限分布</a></li>
<li><a href="#filter">二、主观调整说明</a></li>
<li><a href="#portfolio">三、ETF 组合建议</a></li>
<li><a href="#analysis">四、分析摘要</a></li>
<li><a href="#agent">五、Agent 完整回复</a></li>
<li><a href="#risk">六、风险提示</a></li>
</ul>
</nav>

<section id="quadrant" class="section">
<h2>一、本期四象限分布</h2>
<div class="chart"><img src="data:image/png;base64,{chart_quadrant}" alt="四象限分布" width="600"></div>
<table>
<tr><th>象限</th><th>行业</th></tr>
<tr><td><strong>黄金配置区</strong></td><td>{escape(data.get('golden_industries', '无'))}</td></tr>
<tr><td><strong>左侧观察区</strong></td><td>{escape(data.get('left_side_industries', '无'))}</td></tr>
<tr><td><strong>高危警示区</strong></td><td>{escape(data.get('danger_industries', '无'))}</td></tr>
<tr><td><strong>垃圾规避区</strong></td><td>{escape(data.get('garbage_industries', '无'))}</td></tr>
</table>
</section>

<section id="filter" class="section">
<h2>二、主观调整说明</h2>
<p><strong>观察池</strong>：出口链 {escape(', '.join(obs.get('export_chain', [])))}；政策链 {escape(', '.join(obs.get('policy_chain', [])))}；防守 {escape(', '.join(obs.get('defensive', [])))}</p>
<p><strong>否决行业</strong>：{escape('；'.join(veto) if veto else '无')}</p>
</section>

<section id="portfolio" class="section">
<h2>三、本期 ETF 组合建议</h2>
<div class="chart"><img src="data:image/png;base64,{chart_pie}" alt="组合权重" width="450"></div>
<table>
<tr><th>行业</th><th>代码</th><th>权重</th><th>理由</th></tr>
{etf_rows_html if etf_rows_html else '<tr><td colspan="4">无</td></tr>'}
</table>
<p><strong>现金留存</strong>：{escape(str(data.get('cash_reserve', '10')))}%</p>
</section>

<section id="analysis" class="section">
<h2>四、分析摘要</h2>
<p class="reasoning"><strong>推理链</strong>：{escape(str(data.get('reasoning_chain', '-')))}</p>
<p><strong>新闻交叉验证</strong></p>
<ul>{news_html}</ul>
</section>

<section id="agent" class="section">
<h2>五、Agent 完整回复</h2>
<div class="agent-response">{escape(agent_response)}</div>
</section>

<section id="risk" class="section">
<h2>六、风险提示</h2>
<p><strong>集中度</strong>：{escape(str(risk.get('concentration_risk', '-')))}</p>
<p><strong>流动性</strong>：{escape(str(risk.get('liquidity_risk', '-')))}</p>
<p><strong>宏观风险</strong></p><ul>{risk_macro or '<li>无</li>'}</ul>
<p><strong>行业风险</strong></p><ul>{risk_sector or '<li>无</li>'}</ul>
</section>

<hr>
<p style="color:#7f8c8d;font-size:0.9em">* 本报告由 AI Quant Assistant 自动生成，须经投研/合规审批后方可对客。</p>
</div>
</body>
</html>"""


def _export_docx(data: dict, agent_response: str, chart_quadrant: str, chart_pie: str, out_path: str) -> bool:
    """导出为 Word (.docx)"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("提示: 安装 python-docx 以支持 Word 导出: pip install python-docx")
        return False

    risk = data.get("risk_checks", {})
    obs = data.get("observation_pool", {})
    news = data.get("news_validation", {})
    veto = data.get("veto_list", [])

    doc = Document()
    doc.add_heading("ETF 行业轮动周报", 0)
    doc.add_paragraph().add_run(
        f"报告周：{data.get('report_week', 'N/A')}  |  报告日期：{data.get('decision_date', 'N/A')}  |  "
        f"信号截止：{data.get('decision_date', 'N/A')}  |  因子计算区间：{data.get('data_timestamp', '-')}  |  "
        f"配置版本：{str(data.get('config_version', 'N/A'))}"
    )

    doc.add_heading("一、本期四象限分布", level=1)
    try:
        img_data = base64.b64decode(chart_quadrant)
        doc.add_picture(io.BytesIO(img_data), width=Inches(5.5))
    except Exception:
        doc.add_paragraph("（图表加载失败）")
    table = doc.add_table(rows=5, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "象限", "行业"
    table.rows[1].cells[0].text, table.rows[1].cells[1].text = "黄金配置区", data.get("golden_industries", "无")
    table.rows[2].cells[0].text, table.rows[2].cells[1].text = "左侧观察区", data.get("left_side_industries", "无")
    table.rows[3].cells[0].text, table.rows[3].cells[1].text = "高危警示区", data.get("danger_industries", "无")
    table.rows[4].cells[0].text, table.rows[4].cells[1].text = "垃圾规避区", data.get("garbage_industries", "无")

    doc.add_heading("二、主观调整说明", level=1)
    doc.add_paragraph(
        f"观察池：出口链 {', '.join(obs.get('export_chain', []))}；"
        f"政策链 {', '.join(obs.get('policy_chain', []))}；"
        f"防守 {', '.join(obs.get('defensive', []))}"
    )
    doc.add_paragraph(f"否决行业：{'；'.join(veto) if veto else '无'}")

    doc.add_heading("三、本期 ETF 组合建议", level=1)
    try:
        img_data = base64.b64decode(chart_pie)
        doc.add_picture(io.BytesIO(img_data), width=Inches(4))
    except Exception:
        doc.add_paragraph("（图表加载失败）")

    rows = [
        [row.get("sector", ""), str(row.get("code", "-")), str(row.get("weight", "-")), str(row.get("rationale", "-"))]
        for row in data.get("etf_rows", [])
    ]
    if rows:
        table = doc.add_table(rows=len(rows) + 1, cols=4)
        table.rows[0].cells[0].text, table.rows[0].cells[1].text, table.rows[0].cells[2].text, table.rows[0].cells[3].text = (
            "行业",
            "代码",
            "权重",
            "理由",
        )
        for idx, row in enumerate(rows, 1):
            table.rows[idx].cells[0].text, table.rows[idx].cells[1].text, table.rows[idx].cells[2].text, table.rows[idx].cells[3].text = row
    else:
        doc.add_paragraph("无")
    doc.add_paragraph(f"现金留存：{data.get('cash_reserve', '10')}%")

    doc.add_heading("四、分析摘要", level=1)
    doc.add_paragraph(f"推理链：{data.get('reasoning_chain', '-')}")
    doc.add_paragraph("新闻交叉验证")
    for key, value in (news or {}).items():
        doc.add_paragraph(f"  • {key}：{value}", style="List Bullet")

    doc.add_heading("五、Agent 完整回复", level=1)
    doc.add_paragraph(agent_response)

    doc.add_heading("六、风险提示", level=1)
    doc.add_paragraph(f"集中度：{risk.get('concentration_risk', '-')}")
    doc.add_paragraph(f"流动性：{risk.get('liquidity_risk', '-')}")
    doc.add_paragraph("宏观风险")
    for item in risk.get("macro_risks", []):
        doc.add_paragraph(f"  • {item}", style="List Bullet")
    doc.add_paragraph("行业风险")
    for item in risk.get("sector_risks", []):
        doc.add_paragraph(f"  • {item}", style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("* 本报告由 AI Quant Assistant 自动生成，须经投研/合规审批后方可对客。")
    doc.save(out_path)
    return True


def _export_pdf(html: str, out_path: str) -> bool:
    """将 HTML 转为 PDF（使用 Playwright Chromium，对中文支持好）"""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("提示: 安装 playwright 后执行 'playwright install chromium' 以支持 PDF 导出")
        return False

    try:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as file_obj:
            file_obj.write(html)
            tmp_path = file_obj.name
        try:
            with sync_playwright() as playwright:
                browser = playwright.chromium.launch()
                page = browser.new_page()
                page.goto("file:///" + os.path.abspath(tmp_path).replace("\\", "/"))
                page.pdf(path=out_path, format="A4", print_background=True)
                browser.close()
        finally:
            os.unlink(tmp_path)
        return True
    except Exception as exc:
        print(f"PDF 导出失败: {exc}")
        print("  请确保已执行: playwright install chromium")
        return False


def main(argv=None):
    """argv: 可选，用于被 main.py 等调用时传入空列表，避免 --report 等参数冲突"""
    parser = argparse.ArgumentParser(description="生成 ETF 行业轮动周报")
    parser.add_argument("trace_path", nargs="?", help="trace JSON 文件路径，省略则使用今日最新 trace")
    parser.add_argument(
        "--format",
        "-f",
        choices=["html", "docx", "pdf", "all"],
        default="all",
        help="导出格式: html / docx / pdf / all (默认 all)",
    )
    args = parser.parse_args(argv)

    try:
        trace, trace_path, trace_dir = load_trace(args.trace_path)
    except FileNotFoundError:
        print("未找到今日 trace，请指定路径: python scripts/generate_report_html.py <trace.json>")
        raise SystemExit(1)

    if not args.trace_path:
        print(f"使用 trace: {trace_path}")

    agent_response = load_agent_response(trace_dir)
    data = build_report_data(trace)
    chart_quadrant = _make_quadrant_chart(data)
    chart_pie = _make_weights_pie(data)
    html = _build_html(data, agent_response, chart_quadrant, chart_pie)

    formats = ["html", "docx", "pdf"] if args.format == "all" else [args.format]
    outputs = []

    if "html" in formats:
        out_path = os.path.join(trace_dir, "weekly_report.html")
        with open(out_path, "w", encoding="utf-8") as file_obj:
            file_obj.write(html)
        outputs.append(out_path)

    if "docx" in formats:
        out_path = os.path.join(trace_dir, "weekly_report.docx")
        if _export_docx(data, agent_response, chart_quadrant, chart_pie, out_path):
            outputs.append(out_path)

    if "pdf" in formats:
        out_path = os.path.join(trace_dir, "weekly_report.pdf")
        if _export_pdf(html, out_path):
            outputs.append(out_path)

    print(f"\n报告周（上周）: {data.get('report_week', 'N/A')}")
    for path in outputs:
        print(f"  已生成: {path}")
    if "html" in formats and outputs:
        print("  可在浏览器中打开 HTML 查看图表与完整内容")
    if not outputs:
        print("  未成功生成任何输出")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
